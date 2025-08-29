"""
ææ¡ˆç”Ÿæˆ API è·¯ç”±
================

è™•ç†ç ”ç©¶ææ¡ˆçš„ç”Ÿæˆã€ç·¨è¼¯å’Œä¿®è¨‚åŠŸèƒ½
"""

from fastapi import APIRouter, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import asyncio
import sys
import os
import tempfile
import io
import requests
from docx import Document as DocxDocument
from docx.shared import Inches
from io import BytesIO
import logging
from typing import List, Dict, Any, Optional
from fastapi import APIRouter, HTTPException, UploadFile, File, Form

# é…ç½® logger
logger = logging.getLogger(__name__)
from fastapi.responses import FileResponse
import tempfile
import os
import re
from docx import Document as DocxDocument
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

# æ·»åŠ åŸé …ç›®è·¯å¾‘åˆ° sys.path
sys.path.append(os.path.join(os.path.dirname(__file__), '../../../app'))

# ä¿®æ­£å°å…¥æ–¹å¼
try:
    from backend.services.pubchem_service import chemical_metadata_extractor
except ImportError:
    # å¦‚æœç›´æ¥å°å…¥å¤±æ•—ï¼Œå˜—è©¦ä½¿ç”¨å®Œæ•´è·¯å¾‘ (å·²é‡çµ„ï¼Œä¸å†éœ€è¦)
    # sys.path.append(os.path.join(os.path.dirname(__file__), '../../../'))
    from backend.services.pubchem_service import chemical_metadata_extractor
from backend.services.chemical_service import chemical_service
from langchain_core.documents import Document

# SVG è½‰æ›ä¾è³´æª¢æŸ¥
try:
    from svglib.svglib import svg2rlg
    from reportlab.graphics import renderPDF
    SVGLIB_AVAILABLE = True
except ImportError as e:
    SVGLIB_AVAILABLE = False

# PyMuPDF ç”¨æ–¼ PDF åˆ° PNG è½‰æ›
try:
    import fitz
    PYMUPDF_AVAILABLE = True
except ImportError as e:
    PYMUPDF_AVAILABLE = False

import time
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from PIL import Image

router = APIRouter()

class ProposalRequest(BaseModel):
    """ææ¡ˆç”Ÿæˆè«‹æ±‚æ¨¡å‹"""
    research_goal: str
    user_feedback: Optional[str] = None
    previous_proposal: Optional[str] = None
    retrieval_count: Optional[int] = 10  # é è¨­æª¢ç´¢ 10 å€‹æ–‡æª”

class ProposalResponse(BaseModel):
    """ææ¡ˆç”ŸæˆéŸ¿æ‡‰æ¨¡å‹"""
    proposal: str
    chemicals: List[Dict[str, Any]]
    experiment_detail: Optional[str] = None
    citations: List[Dict[str, str]]
    not_found: List[str]
    # ä»¥å¯åºåˆ—åŒ–çš„çµæ§‹å›å‚³ chunksï¼š[{ page_content, metadata }]
    chunks: List[Dict[str, Any]]
    used_model: Optional[str] = None
    structured_proposal: Optional[Dict[str, Any]] = None


class ProposalRevisionRequest(BaseModel):
    """ææ¡ˆä¿®è¨‚è«‹æ±‚æ¨¡å‹"""
    original_proposal: str
    user_feedback: str
    # ä¾†è‡ªå‰ç«¯çš„å¯åºåˆ—åŒ– chunks
    chunks: List[Dict[str, Any]]
    # å¯é¸çš„æª¢ç´¢åƒæ•¸
    k_new_chunks: Optional[int] = 3  # æ–°chunksæª¢ç´¢æ•¸é‡ï¼Œé è¨­3ï¼ˆé™ä½æŸ¥è©¢é‡ï¼‰


def _serialize_chunks(chunks: List[Any]) -> List[Dict[str, Any]]:
    """å°‡ LangChain Document ç‰©ä»¶åºåˆ—åŒ–ç‚ºå¯å›å‚³çš„ dict çµæ§‹ã€‚"""
    serialized: List[Dict[str, Any]] = []
    for doc in chunks or []:
        try:
            serialized.append({
                "page_content": getattr(doc, "page_content", ""),
                "metadata": getattr(doc, "metadata", {}) or {}
            })
        except Exception:
            continue
    return serialized


def _deserialize_chunks(chunks_like: List[Dict[str, Any]]) -> List[Document]:
    """å°‡å‰ç«¯å‚³ä¾†çš„ dict çµæ§‹é‚„åŸç‚º LangChain Documentã€‚"""
    documents: List[Document] = []
    for item in chunks_like or []:
        page_content = item.get("page_content", "")
        metadata = item.get("metadata", {}) or {}
        documents.append(Document(page_content=page_content, metadata=metadata))
    return documents

@router.post("/proposal/generate", response_model=ProposalResponse)
async def generate_proposal(request: ProposalRequest):
    """
    ç”Ÿæˆç ”ç©¶ææ¡ˆ
    
    Args:
        request: åŒ…å«ç ”ç©¶ç›®æ¨™çš„è«‹æ±‚
        
    Returns:
        ç”Ÿæˆçš„ææ¡ˆå…§å®¹ï¼ŒåŒ…æ‹¬åŒ–å­¸å“ä¿¡æ¯å’Œå¯¦é©—ç´°ç¯€
    """
    import time
    import uuid
    
    # ç”Ÿæˆå”¯ä¸€çš„è«‹æ±‚ ID
    request_id = str(uuid.uuid4())[:8]
    start_time = time.time()
    
    print(f"ğŸš€ [DEBUG-{request_id}] ========== é–‹å§‹è™•ç†ææ¡ˆç”Ÿæˆè«‹æ±‚ ==========")
    print(f"ğŸš€ [DEBUG-{request_id}] æ™‚é–“æˆ³: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"ğŸš€ [DEBUG-{request_id}] æ”¶åˆ°è«‹æ±‚ research_goal = '{request.research_goal}'")
    print(f"ğŸš€ [DEBUG-{request_id}] retrieval_count = {request.retrieval_count}")
    print(f"ğŸš€ [DEBUG-{request_id}] è«‹æ±‚ä¾†æº: {request}")
    
    try:
        print(f"ğŸ” [DEBUG-{request_id}] æº–å‚™èª¿ç”¨ agent_answer with mode='make proposal'")
        
        # å»¶é²å°å…¥ä»¥é¿å…å¾ªç’°å°å…¥å•é¡Œ
        from backend.services.knowledge_service import agent_answer
        
        # èˆ‡ Streamlit Tab1 å°é½Šï¼šä½¿ç”¨æ¨¡å¼ make proposal ç”Ÿæˆææ¡ˆ
        result = agent_answer(request.research_goal, mode="make proposal", k=request.retrieval_count)
        
        print(f"ğŸ” [DEBUG-{request_id}] agent_answer èª¿ç”¨æˆåŠŸ")
        print(f"ğŸ” [DEBUG-{request_id}] result é¡å‹: {type(result)}")
        print(f"ğŸ” [DEBUG-{request_id}] result éµ: {list(result.keys())}")
        print(f"ğŸ” [DEBUG-{request_id}] result['answer'] é•·åº¦: {len(result.get('answer', ''))}")
        print(f"ğŸ” [DEBUG-{request_id}] result['answer'] å…§å®¹: {result.get('answer', '')[:200]}...")

        # å¾å›ç­”ä¸­æŠ½å–åŒ–å­¸å“è³‡è¨Šèˆ‡ææ¡ˆæ­£æ–‡ï¼ˆåŒ…å« SMILES ç¹ªè£½çš„çµæ§‹åœ–ï¼‰
        print(f"ğŸ” [DEBUG-{request_id}] æº–å‚™èª¿ç”¨åŒ–å­¸æœå‹™æå–åŒ–å­¸å“ä¸¦æ·»åŠ çµæ§‹åœ–")
        
        # æª¢æŸ¥æ˜¯å¦æœ‰çµæ§‹åŒ–æ•¸æ“šä¸­çš„ææ–™åˆ—è¡¨
        structured_proposal = result.get("structured_proposal")
        if structured_proposal and structured_proposal.get('materials_list'):
            print(f"ğŸ” [DEBUG-{request_id}] ä½¿ç”¨çµæ§‹åŒ–æ•¸æ“šä¸­çš„ææ–™åˆ—è¡¨: {structured_proposal['materials_list']}")
            # ç›´æ¥ä½¿ç”¨çµæ§‹åŒ–æ•¸æ“šä¸­çš„ææ–™åˆ—è¡¨
            from backend.services.pubchem_service import extract_and_fetch_chemicals, remove_json_chemical_block
            chemical_metadata_list, not_found_list = extract_and_fetch_chemicals(structured_proposal['materials_list'])
            # æ¸…ç†æ–‡æœ¬ä¸­çš„ JSON åŒ–å­¸å“å¡Š
            proposal_answer = remove_json_chemical_block(result.get("answer", ""))
            
            # âœ… ä¿®å¾©ï¼šç‚ºåŒ–å­¸å“æ·»åŠ SMILESç¹ªè£½çš„çµæ§‹åœ–
            print(f"ğŸ” [DEBUG-{request_id}] ç‚ºçµæ§‹åŒ–æ•¸æ“šçš„åŒ–å­¸å“æ·»åŠ SMILESç¹ªè£½")
            print(f"ğŸ” [DEBUG-{request_id}] åŒ–å­¸å“æ•¸é‡: {len(chemical_metadata_list)}")
            
            # æ¸¬è©¦ SMILES-Drawer æ˜¯å¦æ­£å¸¸å·¥ä½œ
            try:
                from backend.services.smiles_drawer import smiles_drawer
                test_smiles = "CCO"  # ä¹™é†‡
                print(f"ğŸ” [DEBUG-{request_id}] æ¸¬è©¦ SMILES-Drawer åŠŸèƒ½...")
                test_svg = smiles_drawer.smiles_to_svg(test_smiles)
                test_png = smiles_drawer.smiles_to_png_base64(test_smiles)
                print(f"ğŸ” [DEBUG-{request_id}] æ¸¬è©¦çµæœ - SVG: {test_svg is not None}, PNG: {test_png is not None}")
            except Exception as e:
                print(f"âŒ [DEBUG-{request_id}] SMILES-Drawer æ¸¬è©¦å¤±æ•—: {e}")
            
            enhanced_chemicals = []
            for i, chemical in enumerate(chemical_metadata_list):
                print(f"ğŸ” [DEBUG-{request_id}] è™•ç†åŒ–å­¸å“ {i+1}/{len(chemical_metadata_list)}: {chemical.get('name', 'Unknown')}")
                print(f"ğŸ” [DEBUG-{request_id}] åŒ–å­¸å“æ•¸æ“šéµ: {list(chemical.keys())}")
                print(f"ğŸ” [DEBUG-{request_id}] SMILES: {chemical.get('smiles', 'N/A')}")
                enhanced_chemical = chemical_service.add_smiles_drawing(chemical)
                enhanced_chemicals.append(enhanced_chemical)
                print(f"ğŸ” [DEBUG-{request_id}] è™•ç†å®Œæˆï¼Œæœ€çµ‚æ•¸æ“šéµ: {list(enhanced_chemical.keys())}")
            chemical_metadata_list = enhanced_chemicals
        else:
            # å›é€€åˆ°å¾æ–‡æœ¬ä¸­æå–
            print(f"ğŸ” [DEBUG-{request_id}] å›é€€åˆ°å¾æ–‡æœ¬ä¸­æå–ææ–™åˆ—è¡¨")
            chemical_metadata_list, not_found_list, proposal_answer = chemical_service.extract_chemicals_with_drawings(
                result.get("answer", "")
            )
        
        print(f"ğŸ” [DEBUG-{request_id}] åŒ–å­¸å“æå–å’Œçµæ§‹åœ–ç”Ÿæˆå®Œæˆ")
        print(f"ğŸ” [DEBUG-{request_id}] proposal_answer é•·åº¦: {len(proposal_answer)}")
        print(f"ğŸ” [DEBUG-{request_id}] chemical_metadata_list æ•¸é‡: {len(chemical_metadata_list)}")

        citations = result.get("citations", [])
        chunks = result.get("chunks", [])
        used_model = result.get("used_model", "unknown")

        # ä¿®å¾© citations ä¸­çš„ page æ¬„ä½é¡å‹å•é¡Œ
        fixed_citations = []
        for citation in citations:
            fixed_citation = citation.copy()
            # ç¢ºä¿ page æ¬„ä½æ˜¯å­—ä¸²
            if "page" in fixed_citation:
                fixed_citation["page"] = str(fixed_citation["page"])
            fixed_citations.append(fixed_citation)

        end_time = time.time()
        duration = end_time - start_time
        
        print(f"âœ… [DEBUG-{request_id}] ========== ææ¡ˆç”Ÿæˆå®Œæˆ ==========")
        print(f"âœ… [DEBUG-{request_id}] ç¸½è€—æ™‚: {duration:.2f} ç§’")
        print(f"âœ… [DEBUG-{request_id}] æª¢ç´¢åˆ°çš„æ–‡æª”æ•¸é‡: {len(chunks)}")
        print(f"âœ… [DEBUG-{request_id}] å¼•ç”¨æ•¸é‡: {len(fixed_citations)}")
        print(f"âœ… [DEBUG-{request_id}] åŒ–å­¸å“æ•¸é‡: {len(chemical_metadata_list)}")

        return ProposalResponse(
            proposal=proposal_answer,
            chemicals=chemical_metadata_list,
            citations=fixed_citations,
            not_found=not_found_list,
            chunks=_serialize_chunks(chunks),
            used_model=used_model,
            structured_proposal=result.get("structured_proposal")
        )
        
    except Exception as e:
        end_time = time.time()
        duration = end_time - start_time
        print(f"âŒ [DEBUG-{request_id}] ========== ææ¡ˆç”Ÿæˆå¤±æ•— ==========")
        print(f"âŒ [DEBUG-{request_id}] ç¸½è€—æ™‚: {duration:.2f} ç§’")
        print(f"âŒ [DEBUG-{request_id}] éŒ¯èª¤: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"ææ¡ˆç”Ÿæˆå¤±æ•—: {str(e)}")

@router.post("/proposal/revise", response_model=ProposalResponse)
async def revise_proposal(request: ProposalRevisionRequest):
    """
    æ ¹æ“šç”¨æˆ¶åé¥‹ä¿®è¨‚ææ¡ˆ
    
    Args:
        request: åŒ…å«åŸå§‹ææ¡ˆå’Œç”¨æˆ¶åé¥‹çš„è«‹æ±‚
        
    Returns:
        ä¿®è¨‚å¾Œçš„ææ¡ˆå…§å®¹
    """
    try:
        # å»¶é²å°å…¥ä»¥é¿å…å¾ªç’°å°å…¥å•é¡Œ
        from backend.services.knowledge_service import agent_answer
        
        # æª¢æŸ¥é–‹ç™¼æ¨¡å¼ç‹€æ…‹
        from backend.core.settings_manager import settings_manager
        is_dev_mode = settings_manager.get_dev_mode_status()
        
        # æ ¹æ“šé–‹ç™¼æ¨¡å¼æ±ºå®šæª¢ç´¢åƒæ•¸
        k_new_chunks = 1 if is_dev_mode else (request.k_new_chunks or 3)
        
        # èˆ‡ Streamlit Tab1 å°é½Šï¼šæ¡ç”¨ generate new idea æ¨¡å¼ï¼Œä¸¦å¸¶å…¥åŸå§‹ææ¡ˆèˆ‡ chunks
        result = agent_answer(
            request.user_feedback,
            mode="generate new idea",
            old_chunks=_deserialize_chunks(request.chunks),
            proposal=request.original_proposal,
            k_new_chunks=k_new_chunks,  # å‚³éæª¢ç´¢åƒæ•¸
        )

        # æª¢æŸ¥æ˜¯å¦æœ‰ç›´æ¥çš„ææ–™åˆ—è¡¨ï¼ˆä¾†è‡ªçµæ§‹åŒ–è¼¸å‡ºï¼‰
        if result.get("materials_list"):
            print(f"ğŸ” [DEBUG] ä½¿ç”¨çµæ§‹åŒ–æ•¸æ“šä¸­çš„ææ–™åˆ—è¡¨: {result['materials_list']}")
            # ç›´æ¥ä½¿ç”¨çµæ§‹åŒ–æ•¸æ“šä¸­çš„ææ–™åˆ—è¡¨
            from backend.services.pubchem_service import extract_and_fetch_chemicals, remove_json_chemical_block
            chemical_metadata_list, not_found_list = extract_and_fetch_chemicals(result["materials_list"])
            # æ¸…ç†æ–‡æœ¬ä¸­çš„ JSON åŒ–å­¸å“å¡Š
            proposal_answer = remove_json_chemical_block(result.get("answer", ""))
        else:
            # å›é€€åˆ°å¾æ–‡æœ¬ä¸­æå–
            print(f"ğŸ” [DEBUG] å›é€€åˆ°å¾æ–‡æœ¬ä¸­æå–ææ–™åˆ—è¡¨")
            chemical_metadata_list, not_found_list, proposal_answer = chemical_metadata_extractor(
                result.get("answer", "")
            )

        # âœ… ä¿®å¾©ï¼šç‚ºåŒ–å­¸å“æ·»åŠ SMILESç¹ªè£½çš„çµæ§‹åœ–
        print(f"ğŸ” [DEBUG] ç‚ºä¿®è¨‚ææ¡ˆçš„åŒ–å­¸å“æ·»åŠ SMILESç¹ªè£½")
        from backend.services.chemical_service import chemical_service
        enhanced_chemicals = []
        for chemical in chemical_metadata_list:
            enhanced_chemical = chemical_service.add_smiles_drawing(chemical)
            enhanced_chemicals.append(enhanced_chemical)
        chemical_metadata_list = enhanced_chemicals

        # ä¿®å¾© citations ä¸­çš„ page æ¬„ä½é¡å‹å•é¡Œ
        fixed_citations = []
        for citation in result.get("citations", []):
            fixed_citation = citation.copy()
            # ç¢ºä¿ page æ¬„ä½æ˜¯å­—ä¸²
            if "page" in fixed_citation:
                fixed_citation["page"] = str(fixed_citation["page"])
            fixed_citations.append(fixed_citation)

        return ProposalResponse(
            proposal=proposal_answer,
            chemicals=chemical_metadata_list,
            citations=fixed_citations,
            not_found=not_found_list,
            chunks=_serialize_chunks(result.get("chunks", [])),
            structured_proposal=result.get("structured_proposal")
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"ææ¡ˆä¿®è¨‚å¤±æ•—: {str(e)}")

class ExperimentDetailRequest(BaseModel):
    """ç”Ÿæˆå¯¦é©—ç´°ç¯€è«‹æ±‚æ¨¡å‹"""
    proposal: str
    chunks: List[Dict[str, Any]]


@router.post("/proposal/experiment-detail")
async def generate_experiment_detail(request: ExperimentDetailRequest):
    """
    ç”Ÿæˆå¯¦é©—ç´°ç¯€
    
    Args:
        proposal: ææ¡ˆå…§å®¹
        chunks: ç›¸é—œæ–‡æª”ç‰‡æ®µ
        
    Returns:
        å¯¦é©—ç´°ç¯€å…§å®¹
    """
    try:
        # å»¶é²å°å…¥ä»¥é¿å…å¾ªç’°å°å…¥å•é¡Œ
        from backend.services.knowledge_service import agent_answer
        
        # èˆ‡ Streamlit Tab1 å°é½Šï¼šç”± agent ä»¥æŒ‡å®šæ¨¡å¼å±•é–‹å¯¦é©—ç´°ç¯€
        result = agent_answer(
            "",
            mode="expand to experiment detail",
            chunks=_deserialize_chunks(request.chunks),
            proposal=request.proposal,
        )

        # âœ… ä¿®å¾©ï¼šæ·»åŠ citationså­—æ®µåˆ°è¿”å›çµæœ
        # å¾resultä¸­ç²å–citationsï¼Œå¦‚æœæ²’æœ‰å‰‡è¿”å›ç©ºåˆ—è¡¨
        citations = result.get("citations", [])
        print(f"ğŸ” [DEBUG] generate_experiment_detail è¿”å›çš„citationsæ•¸é‡: {len(citations)}")
        
        return {
            "experiment_detail": result.get("answer", ""),
            "structured_experiment": result.get("structured_experiment", {}),
            "citations": citations,  # âœ… ä¿®å¾©ï¼šæ·»åŠ citationså­—æ®µ
            "success": True,
            "retry_info": {
                "retry_count": getattr(result, 'retry_count', 0),
                "final_tokens": getattr(result, 'final_tokens', 0)
            }
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"å¯¦é©—ç´°ç¯€ç”Ÿæˆå¤±æ•—: {str(e)}")

@router.get("/proposal/status/{task_id}")
async def get_proposal_status(task_id: str):
    """
    ç²å–ææ¡ˆç”Ÿæˆä»»å‹™ç‹€æ…‹
    
    Args:
        task_id: ä»»å‹™ ID
        
    Returns:
        ä»»å‹™ç‹€æ…‹ä¿¡æ¯
    """
    # TODO: å¯¦ç¾ä»»å‹™ç‹€æ…‹è¿½è¹¤
    return {
        "task_id": task_id,
        "status": "completed",
        "progress": 100
    }

def clean_text_for_xml(text):
    """æ¸…ç†æ–‡æœ¬ä»¥ç¢ºä¿XMLå…¼å®¹æ€§"""
    if not text:
        return ""
    # ç§»é™¤NULLå­—ç¯€å’Œæ§åˆ¶å­—ç¬¦
    cleaned = "".join(char for char in text if ord(char) >= 32 or char in '\n\r\t')
    # ç§»é™¤å…¶ä»–å¯èƒ½å°è‡´XMLå•é¡Œçš„å­—ç¬¦
    cleaned = cleaned.replace('\x00', '')  # NULLå­—ç¯€
    cleaned = cleaned.replace('\x01', '')  # SOH
    cleaned = cleaned.replace('\x02', '')  # STX
    cleaned = cleaned.replace('\x03', '')  # ETX
    cleaned = cleaned.replace('\x04', '')  # EOT
    cleaned = cleaned.replace('\x05', '')  # ENQ
    cleaned = cleaned.replace('\x06', '')  # ACK
    cleaned = cleaned.replace('\x07', '')  # BEL
    cleaned = cleaned.replace('\x08', '')  # BS
    cleaned = cleaned.replace('\x0b', '')  # VT
    cleaned = cleaned.replace('\x0c', '')  # FF
    cleaned = cleaned.replace('\x0e', '')  # SO
    cleaned = cleaned.replace('\x0f', '')  # SI
    return cleaned

def clean_markdown_text(text):
    """æ¸…ç† markdown æ ¼å¼ï¼Œè½‰æ›ç‚ºç´”æ–‡æœ¬"""
    if not text:
        return ""
    
    import re
    
    # ç§»é™¤ markdown æ ¼å¼
    cleaned = text
    cleaned = re.sub(r'\*\*(.*?)\*\*', r'\1', cleaned)  # ç§»é™¤ç²—é«”æ¨™è¨˜
    cleaned = re.sub(r'\*(.*?)\*', r'\1', cleaned)  # ç§»é™¤æ–œé«”æ¨™è¨˜
    cleaned = re.sub(r'`(.*?)`', r'\1', cleaned)  # ç§»é™¤ä»£ç¢¼æ¨™è¨˜
    cleaned = re.sub(r'^#+\s*(.*)$', r'\1', cleaned, flags=re.MULTILINE)  # ç§»é™¤æ¨™é¡Œæ¨™è¨˜
    cleaned = re.sub(r'^\s*[-*+]\s+', '- ', cleaned, flags=re.MULTILINE)  # çµ±ä¸€é …ç›®ç¬¦è™Ÿ
    cleaned = re.sub(r'^\s*\d+\.\s+', '', cleaned, flags=re.MULTILINE)  # ç§»é™¤ç·¨è™Ÿ
    cleaned = re.sub(r'\n\s*\n\s*\n', '\n\n', cleaned)  # ç§»é™¤å¤šé¤˜ç©ºè¡Œ
    cleaned = re.sub(r'\n\s*\*\*', '\n', cleaned)  # ç§»é™¤ç²—é«”å‰çš„æ›è¡Œ
    cleaned = re.sub(r'\*\*\s*\n', '\n', cleaned)  # ç§»é™¤ç²—é«”å¾Œçš„æ›è¡Œ
    
    return cleaned

def get_nfpa_icon_image_stream(nfpa_code: str) -> io.BytesIO:
    """
    ç”¨ Headless Chrome æŠ“å– PubChem NFPA åœ–ï¼Œå›å‚³ BytesIO PNG åœ–ç‰‡
    """
    nfpa_svg_url = f"https://pubchem.ncbi.nlm.nih.gov/image/nfpa.cgi?code={nfpa_code}"

    chrome_options = Options()
    chrome_options.add_argument("--headless=new")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--window-size=300,300")
    chrome_options.add_argument("--hide-scrollbars")

    driver = webdriver.Chrome(options=chrome_options)
    driver.get(nfpa_svg_url)
    time.sleep(1)

    png_data = driver.get_screenshot_as_png()
    driver.quit()

    image = Image.open(io.BytesIO(png_data))
    cropped = image.crop((0, 0, 100, 100))  # ä½ å·²ç¶“æ¸¬è©¦å¥½çš„ç¯„åœ
    output = io.BytesIO()
    cropped.save(output, format="PNG")
    output.seek(0)
    return output

class DocxRequest(BaseModel):
    """DOCX ç”Ÿæˆè«‹æ±‚æ¨¡å‹"""
    proposal: str
    chemicals: List[Dict[str, Any]]
    not_found: List[str]
    experiment_detail: Optional[str] = ""
    citations: List[Dict[str, Any]]  # âœ… ä¿®å¾©ï¼šæ”¹ç‚º Any ä»¥æ”¯æŒæ•¸å­—é¡å‹çš„ page å­—æ®µ

@router.post("/proposal/generate-docx")
async def generate_docx(request: DocxRequest):
    """
    ç”Ÿæˆ DOCX æ–‡ä»¶
    
    Args:
        request: åŒ…å«æ‰€æœ‰ææ¡ˆæ•¸æ“šçš„è«‹æ±‚
        
    Returns:
        DOCX æ–‡ä»¶ä¸‹è¼‰éŸ¿æ‡‰
    """
    tmp_path = None
    try:
        print(f"ğŸ” BACKEND DEBUG: é–‹å§‹ç”Ÿæˆ DOCX æ–‡ä»¶")
        print(f"ğŸ” BACKEND DEBUG: proposal é•·åº¦: {len(request.proposal)}")
        print(f"ğŸ” BACKEND DEBUG: chemicals æ•¸é‡: {len(request.chemicals)}")
        print(f"ğŸ” BACKEND DEBUG: experiment_detail é•·åº¦: {len(request.experiment_detail)}")
        print(f"ğŸ” BACKEND DEBUG: citations æ•¸é‡: {len(request.citations)}")
        
        doc = DocxDocument()
        doc.add_heading("AI Generated Research Proposal", 0)

        # Proposal Section
        doc.add_heading("Proposal", level=1)
        proposal_text = clean_text_for_xml(clean_markdown_text(request.proposal))
        doc.add_paragraph(proposal_text)

        # Chemical Table
        doc.add_heading("Chemical Summary Table", level=1)
        table = doc.add_table(rows=1, cols=8)  # å¤šå…©æ¬„ï¼šStructure + Safety
        hdr = table.rows[0].cells
        hdr[0].text = "Structure"
        hdr[1].text = "Name"
        hdr[2].text = "Formula"
        hdr[3].text = "MW"
        hdr[4].text = "Boiling Point (Â°C)"
        hdr[5].text = "Melting Point (Â°C)"
        hdr[6].text = "CAS No."
        hdr[7].text = "Safety Icons"

        for chem in request.chemicals:
            row = table.add_row().cells

            # å„ªå…ˆä½¿ç”¨ SMILES ç¹ªè£½çš„çµæ§‹åœ–
            if chem.get("png_structure"):
                try:
                    # å¾ Base64 è½‰æ›ç‚ºåœ–ç‰‡æµ
                    import base64
                    img_data = base64.b64decode(chem["png_structure"].split(",")[1])
                    img_stream = BytesIO(img_data)
                    row[0].paragraphs[0].add_run().add_picture(img_stream, width=Inches(1))
                    print(f"âœ… ä½¿ç”¨ SMILES ç¹ªè£½çš„çµæ§‹åœ–: {chem.get('name', 'Unknown')}")
                except Exception as e:
                    print(f"âš ï¸ SMILES åœ–ç‰‡æ’å…¥å¤±æ•—: {chem.get('name', 'Unknown')}, {e}")
                    row[0].text = "SMILES image error"
            elif chem.get("image_url"):
                # å‚™ç”¨æ–¹æ¡ˆï¼šä½¿ç”¨åŸæœ‰çš„ URL åœ–ç‰‡
                try:
                    response = requests.get(chem["image_url"], verify=False, timeout=5)
                    if response.status_code == 200:
                        img_stream = BytesIO(response.content)
                        row[0].paragraphs[0].add_run().add_picture(img_stream, width=Inches(1))
                        print(f"âœ… ä½¿ç”¨ URL åœ–ç‰‡: {chem.get('name', 'Unknown')}")
                    else:
                        row[0].text = "Image not found"
                        print(f"âš ï¸ URL åœ–ç‰‡ä¸‹è¼‰å¤±æ•—: {chem.get('name', 'Unknown')}")
                except Exception as e:
                    print(f"âš ï¸ åœ–ç‰‡ä¸‹è¼‰å¤±æ•—: {chem['image_url']}, {e}")
                    row[0].text = "Image error"
            else:
                row[0].text = "No image"
                print(f"âš ï¸ æ²’æœ‰åœ–ç‰‡è³‡æ–™: {chem.get('name', 'Unknown')}")

            # æ–‡å­—æ¬„ä½ - ä½¿ç”¨æ¸…ç†å‡½æ•¸
            row[1].text = clean_text_for_xml(chem.get("name", "-") or "-")
            row[2].text = clean_text_for_xml(chem.get("formula", "-") or "-")
            row[3].text = clean_text_for_xml(str(chem.get("weight", "-") or "-"))
            row[4].text = clean_text_for_xml(str(chem.get("boiling_point_c", "-") or "-"))
            row[5].text = clean_text_for_xml(str(chem.get("melting_point_c", "-") or "-"))
            row[6].text = clean_text_for_xml(chem.get("cas", "-") or "-")

            # Safety iconsï¼ˆæ”¯æ´å¤šå¼µ GHS + 1 å¼µ NFPAï¼‰
            icons_cell = row[7].paragraphs[0]
            ghs_icons = chem.get("safety_icons", {}).get("ghs_icons", [])
            nfpa_icon_url = chem.get("safety_icons", {}).get("nfpa_image")
            
            for icon_url in ghs_icons:
                try:
                    icon_resp = requests.get(icon_url, verify=False, timeout=5)
                    if icon_resp.status_code == 200:
                        # ä½¿ç”¨ svglib + PyMuPDF è½‰æ› SVG ç‚º PNG
                        if SVGLIB_AVAILABLE and PYMUPDF_AVAILABLE:
                            try:
                                # æ­¥é©Ÿ1ï¼šSVG è½‰ PDF
                                drawing = svg2rlg(io.BytesIO(icon_resp.content))
                                if drawing:
                                    # å‰µå»ºè‡¨æ™‚ PDF æ–‡ä»¶
                                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                                        renderPDF.drawToFile(drawing, tmp_pdf.name)
                                    
                                    # æ­¥é©Ÿ2ï¼šPDF è½‰ PNG
                                    pdf_document = fitz.open(tmp_pdf.name)
                                    page = pdf_document[0]
                                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x ç¸®æ”¾
                                    
                                    # å°‡ PNG è½‰æ›ç‚º BytesIO
                                    png_stream = io.BytesIO()
                                    png_stream.write(pix.tobytes("png"))
                                    png_stream.seek(0)
                                    
                                    # æ’å…¥åˆ° Word æ–‡æª”
                                    run = icons_cell.add_run()
                                    run.add_picture(png_stream, width=Inches(0.3))
                                    
                                    # æ¸…ç†è‡¨æ™‚æ–‡ä»¶
                                    pdf_document.close()
                                    os.unlink(tmp_pdf.name)
                                    
                                else:
                                    print(f"âš ï¸ SVG è½‰æ›å¤±æ•— - drawing ç‚º None: {icon_url}")
                            except Exception as e:
                                print(f"âš ï¸ SVG è½‰æ›éŒ¯èª¤: {icon_url}, {e}")
                        else:
                            print(f"âš ï¸ svglib æˆ– PyMuPDF ä¸å¯ç”¨ï¼Œç„¡æ³•è½‰æ› SVG: {icon_url}")
                except Exception as e:
                    print(f"âš ï¸ Failed to convert or insert icon: {icon_url}, {e}")

            if nfpa_icon_url:
                try:
                    # å¾ URL æ“·å–å‡º NFPA codeï¼Œä¾‹å¦‚ https://...code=130
                    from urllib.parse import urlparse, parse_qs
                    parsed = urlparse(nfpa_icon_url)
                    nfpa_code = parse_qs(parsed.query).get("code", [""])[0]

                    if nfpa_code:
                        image_stream = get_nfpa_icon_image_stream(nfpa_code)
                        if image_stream:
                            run = icons_cell.add_run()
                            run.add_picture(image_stream, width=Inches(0.3))
                    else:
                        print(f"âš ï¸ ç„¡æ³•å¾ URL æ“·å– NFPA code: {nfpa_icon_url}")
                except Exception as e:
                    print(f"âš ï¸ Failed to convert or insert NFPA icon: {nfpa_icon_url}, {e}")

        # Not Found Chemicals
        if request.not_found:
            doc.add_heading("Not Found Chemicals", level=2)
            for name in request.not_found:
                doc.add_paragraph(f"- {clean_text_for_xml(name)}")

        # Experiment Details
        doc.add_heading("Experimental Plan", level=1)
        experiment_text = clean_text_for_xml(clean_markdown_text(request.experiment_detail))
        doc.add_paragraph(experiment_text)

        # Citations
        doc.add_heading("Citations", level=1)
        for i, c in enumerate(request.citations, 1):
            title = clean_text_for_xml(c.get('title', ''))
            page = clean_text_for_xml(str(c.get('page', '')))
            snippet = clean_text_for_xml(c.get('snippet', ''))
            doc.add_paragraph(f"[{i}] {title} | Page {page} | Snippet: {snippet}")

        # Save and Download
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
            logger.info(f"DOCX æ–‡ä»¶å·²ä¿å­˜åˆ°: {tmp_path}")

        logger.info("æº–å‚™è¿”å› FileResponse")
        return FileResponse(
            path=tmp_path,
            filename="proposal_report.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        logger.error(f"DOCX ç”Ÿæˆå¤±æ•—: {str(e)}", exc_info=True)
        # æ¸…ç†è‡¨æ™‚æ–‡ä»¶
        if 'tmp_path' in locals() and tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError as cleanup_error:
                logger.warning(f"æ¸…ç†è‡¨æ™‚æ–‡ä»¶å¤±æ•—: {cleanup_error}")
        raise HTTPException(status_code=500, detail=f"DOCX ç”Ÿæˆå¤±æ•—: {str(e)}") 

@router.post("/proposal/test-docx")
async def test_docx_generation():
    """
    æ¸¬è©¦ DOCX ç”ŸæˆåŠŸèƒ½
    """
    try:
        logger.info("é–‹å§‹æ¸¬è©¦ DOCX ç”Ÿæˆ")
        
        doc = DocxDocument()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("This is a test document to verify DOCX generation works.")

        # Save and Download
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
            logger.info(f"æ¸¬è©¦ DOCX æ–‡ä»¶å·²ä¿å­˜åˆ°: {tmp_path}")

        return FileResponse(
            path=tmp_path,
            filename="test_document.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        logger.error(f"æ¸¬è©¦ DOCX ç”Ÿæˆå¤±æ•—: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"æ¸¬è©¦ DOCX ç”Ÿæˆå¤±æ•—: {str(e)}") 