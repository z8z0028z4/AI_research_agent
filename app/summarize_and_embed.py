
from typing import List, Dict
import os
import fitz
import docx
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from config import PAPER_DIR, VECTOR_INDEX_DIR

def embed_documents_from_metadata(metadata_list: List[Dict]):
    embeddings = OpenAIEmbeddings()
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    documents = []

    for meta in metadata_list:
        filename = meta.get("new_filename", "")
        file_path = os.path.join(PAPER_DIR, filename)

        if not os.path.exists(file_path):
            print(f"⚠️ 檔案不存在，略過：{filename}")
            continue

        if filename.endswith(".pdf"):
            doc = fitz.open(file_path)
            for page_index, page in enumerate(doc):
                text = page.get_text()
                if not text.strip():
                    continue
                chunks = text_splitter.split_text(text)
                for i, chunk in enumerate(chunks):
                    documents.append(Document(
                        page_content=chunk,
                        metadata={
                            "tracing_number": meta.get("tracing_number", ""),
                            "title": meta.get("title", ""),
                            "type": meta.get("type", ""),
                            "filename": filename,
                            "page_number": page_index + 1,
                            "chunk_index": i
                        }
                    ))
        elif filename.endswith(".docx"):
            docx_text = "".join([para.text for para in docx.Document(file_path).paragraphs if para.text.strip()])
            chunks = text_splitter.split_text(docx_text)
            for i, chunk in enumerate(chunks):
                documents.append(Document(
                    page_content=chunk,
                    metadata={
                        "tracing_number": meta.get("tracing_number", ""),
                        "title": meta.get("title", ""),
                        "type": meta.get("type", ""),
                        "filename": filename,
                        "page_number": None,
                        "chunk_index": i
                    }
                ))
        else:
            print(f"⚠️ 不支援的格式，略過：{filename}")
            continue

    if documents:
        db = Chroma.from_documents(documents, embeddings, persist_directory=VECTOR_INDEX_DIR)
        db.persist()
        print(f"✅ 嵌入完成，共 {len(documents)} 段，已包含 page_number")
    else:
        print("⚠️ 無資料可嵌入")

