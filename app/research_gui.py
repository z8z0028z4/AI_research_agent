import streamlit as st
from perplexity_search_fallback import ask_perplexity
from search_agent import search_and_download_only
import os
import tempfile
import re
from knowledge_agent import agent_answer
from browser import select_files, select_files_paper_mode
from file_upload import process_uploaded_files
from chunk_embedding import embed_documents_from_metadata, embed_experiment_txt_batch
import pandas as pd
from excel_to_txt_by_row import export_new_experiments_to_txt
from config import EXPERIMENT_DIR
from pubchem_handler import chemical_metadata_extractor
from rag_core import build_detail_experimental_plan_prompt
from docx import Document
from docx.shared import Inches
import io
from io import BytesIO
import requests

# SVG 轉換依賴檢查
try:
    from svglib.svglib import svg2rlg
    from reportlab.graphics import renderPDF
    SVGLIB_AVAILABLE = True
except ImportError as e:
    SVGLIB_AVAILABLE = False

# PyMuPDF 用於 PDF 到 PNG 轉換
try:
    import fitz
    PYMUPDF_AVAILABLE = True
except ImportError as e:
    PYMUPDF_AVAILABLE = False

import time
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from PIL import Image

def update_gui_with_docx_data():
    docx_data = st.session_state.get("docx_data", {})

    st.markdown("### 🤖 Generated proposal")
    st.markdown(docx_data.get("proposal", ""))

    render_chemical_table(docx_data.get("chemicals", []))

    if docx_data.get("not_found"):
        st.markdown("### ⚠️ 以下化學品未能查詢成功")
        for name in docx_data["not_found"]:
            st.markdown(f"- {name}")

    if docx_data.get("experiment_detail"):
        st.markdown("### 🔬 Suggested experiment details")
        st.markdown(docx_data["experiment_detail"])

    with st.expander("📚 Citations ", expanded=False):
        st.markdown("### 📚 Citations ")
        for i, citation in enumerate(docx_data.get("citations", []), start=1):
            title = citation.get("title", "未知")
            page = citation.get("page", "?")
            snippet = citation.get("snippet", "...")
            st.markdown(f"**[{i}]** {title} | 頁碼：{page} | 段落開頭：{snippet}")
def format_references_block(text):
    refs = []
    in_reference = False
    for line in text.splitlines():
        if line.strip().lower().startswith("reference") or line.strip().lower().startswith("references"):
            in_reference = True
            continue
        if in_reference and line.strip():
            refs.append(line.strip())
    markdown_links = []
    for ref in refs:
        match = re.match(r"\[(\d+)\]\s*(.*?)(https?://\S+)", ref)
        if match:
            idx, title, url = match.groups()
            markdown_links.append(f"[{idx}] [{title.strip()}]({url.strip()})")
        else:
            markdown_links.append(ref)
    return markdown_links
def render_chemical_table(chemical_metadata_list):
    st.markdown("### 🧪 Chemical Summary Table")

    for chem in chemical_metadata_list:
        cols = st.columns([2, 2, 3, 3])

        # Chemical Name
        with cols[0]:
            url = chem.get("pubchem_url")
            if url:
                st.markdown(f"[**{chem.get('name', 'Unknown')}**]({url})")
            else:
                st.markdown(f"**{chem.get('name', 'Unknown')}**")
            

        # Structure Image
        with cols[1]:
            st.image(chem.get("image_url", ""), width=200)

        # Properties block
        with cols[2]:
            st.markdown("**Properties**")
            st.markdown(f"- **Formula**: `{chem.get('formula', '-')}`")
            st.markdown(f"- **MW**: `{chem.get('weight', '-')}`")
            st.markdown(f"- **Boiling Point**: `{chem.get('boiling_point_c', '-')}`")
            st.markdown(f"- **Melting Point**: `{chem.get('melting_point_c', '-')}`")
            st.markdown(f"- **CAS No.**: `{chem.get('cas', '-')}`")
            st.markdown(f"- **SMILES**: `{chem.get('smiles', '-')}`")

        # Safety block (GHS + NFPA)
        with cols[3]:
            st.markdown("**Handling Safety**")
            if chem.get("safety_icons", {}).get("nfpa_image"):
                st.image(chem["safety_icons"]["nfpa_image"], width=60)

            ghs = chem.get("safety_icons", {}).get("ghs_icons", [])
            if ghs:
                st.image(ghs, width=50)

        st.markdown("---")
def get_nfpa_icon_image_stream(nfpa_code: str) -> io.BytesIO:
    """
    用 Headless Chrome 抓取 PubChem NFPA 圖，回傳 BytesIO PNG 圖片
    """
    nfpa_svg_url = f"https://pubchem.ncbi.nlm.nih.gov/image/nfpa.cgi?code={nfpa_code}"

    chrome_options = Options()
    chrome_options.add_argument("--headless=new")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--window-size=300,300")
    chrome_options.add_argument("--hide-scrollbars")

    driver = webdriver.Chrome(options=chrome_options)
    driver.get(nfpa_svg_url)
    time.sleep(1)

    png_data = driver.get_screenshot_as_png()
    driver.quit()

    image = Image.open(io.BytesIO(png_data))
    cropped = image.crop((0, 0, 100, 100))  # 你已經測試好的範圍
    output = io.BytesIO()
    cropped.save(output, format="PNG")
    output.seek(0)
    return output

TAB_FLAGS = {
    "tab_1_proposal_generator" : True,
    "tab_2_embedding":True,
    "tab_3_search_pdf": False,
    "tab_4_perplexity_search": False,
    "tab_5_research_assitant": False
}

st.set_page_config(page_title="AI Chemist", layout="wide")
st.title("🧪 AI Chemist")

# 根據 TAB_FLAGS 動態生成 tabs
tab_names = []
tab_keys = []

if TAB_FLAGS.get("tab_1_proposal_generator"):
    tab_names.append("Proposal generator")
    tab_keys.append("tab1")
if TAB_FLAGS.get("tab_2_embedding"):
    tab_names.append("📥 論文/實驗資料上傳")
    tab_keys.append("tab3")
if TAB_FLAGS.get("tab_3_search_pdf"):
    tab_names.append("🔍 搜尋外部文獻並下載 PDF")
    tab_keys.append("tab2")
if TAB_FLAGS.get("tab_4_perplexity_search"):
    tab_names.append("🌐 使用 Perplexity 搜尋")
    tab_keys.append("tab4")
if TAB_FLAGS.get("tab_5_research_assitant"):
    tab_names.append("📘 知識庫助理, archive functions")
    tab_keys.append("tab5")

tabs = st.tabs(tab_names)
tab_dict = dict(zip(tab_keys, tabs))


if TAB_FLAGS["tab_1_proposal_generator"]:
    with tab_dict["tab1"]:
        st.subheader("✍️ Proposal Generator")

        q1 = st.text_area("Please enter your research goal：", height=100, key="search1")

        if st.button("✍️ Generate proposal", key="generate_proposal_btn"):
            with st.spinner("📖 Gathering information from knowledgebase and drafting proposal..."):
                result = agent_answer(q1, mode="make proposal")
                chemical_metadata_list, not_found_list, proposal_answer = chemical_metadata_extractor(result["answer"])

                st.session_state["docx_data"] = {
                    "proposal": proposal_answer,
                    "chemicals": chemical_metadata_list,
                    "not_found": not_found_list,
                    "citations": result.get("citations", []),
                    "experiment_detail": ""
                }
                st.session_state["proposal_chunks"] = result.get("chunks", [])
                st.rerun()

        if "docx_data" in st.session_state:
            update_gui_with_docx_data()
        
            st.markdown("### 💡 Don't like the proposal？ Provide your opinion here")
            user_reason = st.text_input("How you want to revise?", key="revise_reason")

            if st.button("💡 Generate New Idea", key="new_idea_btn"): #修改proposal
                with st.spinner("🔄 Revise proposal based on your reason and knowledgebase..."):
                    old_chunks = st.session_state.get("proposal_chunks", [])
                    past_proposal = st.session_state["docx_data"].get("proposal", "")
                    result = agent_answer(
                        user_reason,
                        mode="generate new idea",
                        chunks=old_chunks,
                        proposal=past_proposal
                    )
                    chemical_metadata_list, not_found_list, proposal_answer = chemical_metadata_extractor(result["answer"])
                    st.session_state["docx_data"].update({
                        "proposal": proposal_answer,
                        "chemicals": chemical_metadata_list,
                        "not_found": not_found_list,
                        "citations": result.get("citations", []),
                        "experiment_detail": ""
                    })
                    st.session_state["proposal_chunks"] = result.get("chunks", [])
                    st.rerun()

            if st.button("✅ Accept & Generate Experiment Detail", key="accept_btn"): #生成實驗細節
                        with st.spinner("🧪 Generating experimental details..."):
                            chunks = st.session_state.get("proposal_chunks", [])
                            proposal = st.session_state["docx_data"].get("proposal", "")
                            result = agent_answer(
                                "",
                                mode="expand to experiment detail",
                                chunks=chunks,
                                proposal=proposal
                            )
                            st.session_state["docx_data"]["experiment_detail"] = result["answer"]
                            st.rerun()
                            
            if st.button("📥 Prepare proposal.docx"): #下載文件按鈕
                with st.spinner("🔄 Preparing your proposal..."):
                    docx_data = st.session_state.get("docx_data", {})
                    doc = Document()
                    doc.add_heading("AI Generated Research Proposal", 0)

                    # 清理文本以確保XML兼容性
                    def clean_text_for_xml(text):
                        """清理文本以確保XML兼容性"""
                        if not text:
                            return ""
                        # 移除NULL字節和控制字符
                        cleaned = "".join(char for char in text if ord(char) >= 32 or char in '\n\r\t')
                        # 移除其他可能導致XML問題的字符
                        cleaned = cleaned.replace('\x00', '')  # NULL字節
                        cleaned = cleaned.replace('\x01', '')  # SOH
                        cleaned = cleaned.replace('\x02', '')  # STX
                        cleaned = cleaned.replace('\x03', '')  # ETX
                        cleaned = cleaned.replace('\x04', '')  # EOT
                        cleaned = cleaned.replace('\x05', '')  # ENQ
                        cleaned = cleaned.replace('\x06', '')  # ACK
                        cleaned = cleaned.replace('\x07', '')  # BEL
                        cleaned = cleaned.replace('\x08', '')  # BS
                        cleaned = cleaned.replace('\x0b', '')  # VT
                        cleaned = cleaned.replace('\x0c', '')  # FF
                        cleaned = cleaned.replace('\x0e', '')  # SO
                        cleaned = cleaned.replace('\x0f', '')  # SI
                        return cleaned

                    # Proposal Section
                    doc.add_heading("Proposal", level=1)
                    proposal_text = clean_text_for_xml(docx_data.get("proposal", ""))
                    doc.add_paragraph(proposal_text)

                    # Chemical Table
                    doc.add_heading("Chemical Summary Table", level=1)
                    table = doc.add_table(rows=1, cols=8)  # 多兩欄：Structure + Safety
                    hdr = table.rows[0].cells
                    hdr[0].text = "Structure"
                    hdr[1].text = "Name"
                    hdr[2].text = "Formula"
                    hdr[3].text = "MW"
                    hdr[4].text = "Boiling Point (°C)"
                    hdr[5].text = "Melting Point (°C)"
                    hdr[6].text = "CAS No."
                    hdr[7].text = "Safety Icons"

                    for chem in docx_data.get("chemicals", []):
                        row = table.add_row().cells

                        # 結構圖片
                        img_url = chem.get("image_url")
                        if img_url:
                            try:
                                response = requests.get(img_url, verify=False, timeout=5)
                                if response.status_code == 200:
                                    img_stream = BytesIO(response.content)
                                    row[0].paragraphs[0].add_run().add_picture(img_stream, width=Inches(1))
                                else:
                                    row[0].text = "Image not found"
                            except:
                                row[0].text = "Image error"
                        else:
                            row[0].text = "No image"

                        # 文字欄位 - 使用清理函數
                        row[1].text = clean_text_for_xml(chem.get("name", "-") or "-")
                        row[2].text = clean_text_for_xml(chem.get("formula", "-") or "-")
                        row[3].text = clean_text_for_xml(str(chem.get("weight", "-") or "-"))
                        row[4].text = clean_text_for_xml(str(chem.get("boiling_point_c", "-") or "-"))
                        row[5].text = clean_text_for_xml(str(chem.get("melting_point_c", "-") or "-"))
                        row[6].text = clean_text_for_xml(chem.get("cas", "-") or "-")

                        # Safety icons（支援多張 GHS + 1 張 NFPA）
                        icons_cell = row[7].paragraphs[0]
                        ghs_icons = chem.get("safety_icons", {}).get("ghs_icons", [])
                        nfpa_icon_url = chem.get("safety_icons", {}).get("nfpa_image")
                        for icon_url in ghs_icons:
                            try:
                                icon_resp = requests.get(icon_url, verify=False, timeout=5)
                                if icon_resp.status_code == 200:
                                    # 使用 svglib + PyMuPDF 轉換 SVG 為 PNG
                                    if SVGLIB_AVAILABLE and PYMUPDF_AVAILABLE:
                                        try:
                                            # 步驟1：SVG 轉 PDF
                                            drawing = svg2rlg(io.BytesIO(icon_resp.content))
                                            if drawing:
                                                # 創建臨時 PDF 文件
                                                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                                                    renderPDF.drawToFile(drawing, tmp_pdf.name)
                                                
                                                # 步驟2：PDF 轉 PNG
                                                pdf_document = fitz.open(tmp_pdf.name)
                                                page = pdf_document[0]
                                                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x 縮放
                                                
                                                # 將 PNG 轉換為 BytesIO
                                                png_stream = io.BytesIO()
                                                png_stream.write(pix.tobytes("png"))
                                                png_stream.seek(0)
                                                
                                                # 插入到 Word 文檔
                                                run = icons_cell.add_run()
                                                run.add_picture(png_stream, width=Inches(0.3))
                                                
                                                # 清理臨時文件
                                                pdf_document.close()
                                                os.unlink(tmp_pdf.name)
                                                
                                            else:
                                                print(f"⚠️ SVG 轉換失敗 - drawing 為 None: {icon_url}")
                                        except Exception as e:
                                            print(f"⚠️ SVG 轉換錯誤: {icon_url}, {e}")
                                    else:
                                        print(f"⚠️ svglib 或 PyMuPDF 不可用，無法轉換 SVG: {icon_url}")
                            except Exception as e:
                                print(f"⚠️ Failed to convert or insert icon: {icon_url}, {e}")

                        

                        if nfpa_icon_url:
                            try:
                                # 從 URL 擷取出 NFPA code，例如 https://...code=130
                                from urllib.parse import urlparse, parse_qs
                                parsed = urlparse(nfpa_icon_url)
                                nfpa_code = parse_qs(parsed.query).get("code", [""])[0]

                                if nfpa_code:
                                    image_stream = get_nfpa_icon_image_stream(nfpa_code)
                                    if image_stream:
                                        run = icons_cell.add_run()
                                        run.add_picture(image_stream, width=Inches(0.3))
                                else:
                                    print(f"⚠️ 無法從 URL 擷取 NFPA code: {nfpa_icon_url}")
                            except Exception as e:
                                print(f"⚠️ Failed to convert or insert NFPA icon: {nfpa_icon_url}, {e}")
                            
                            


                    # Not Found Chemicals
                    not_found = docx_data.get("not_found", [])
                    if not_found:
                        doc.add_heading("Not Found Chemicals", level=2)
                        for name in not_found:
                            doc.add_paragraph(f"- {clean_text_for_xml(name)}")

                    # Experiment Details
                    doc.add_heading("Experimental Plan", level=1)
                    experiment_text = clean_text_for_xml(docx_data.get("experiment_detail", ""))
                    doc.add_paragraph(experiment_text)

                    # Citations
                    doc.add_heading("Citations", level=1)
                    for i, c in enumerate(docx_data.get("citations", []), 1):
                        title = clean_text_for_xml(c.get('title', ''))
                        page = clean_text_for_xml(str(c.get('page', '')))
                        snippet = clean_text_for_xml(c.get('snippet', ''))
                        doc.add_paragraph(f"[{i}] {title} | Page {page} | Snippet: {snippet}")

                    # Save and Download
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                        doc.save(tmp.name)
                        tmp_path = tmp.name

                    with open(tmp_path, "rb") as f:
                        st.download_button(
                            label="📥 Click to download",
                            data=f,
                            file_name="proposal_report.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

                    os.unlink(tmp_path)
        


if TAB_FLAGS["tab_2_embedding"]:
    with tab_dict["tab3"]:
        if "processed_files" not in st.session_state:
            st.session_state.processed_files = set()


        file_info = select_files() 
        if file_info:
            file_type = file_info["type"] 
            
            st.markdown(f"🔎 資料類型：{file_type}")
            for f in file_info["files"]:
                st.write(f"📄 {f.name}")

            if file_type == "📄 論文資料":
                new_files = [f for f in file_info["files"] if f.name not in st.session_state.processed_files]

                if not new_files:
                    st.info("✅ 所有檔案都已處理過，不重複處理。")
                else:
                    with st.spinner("📥 處理論文資料中..."):
                        new_file_paths = []
                        messages = []

                        def update_status(msg):
                            messages.append(msg)

                        for f in new_files:
                            suffix = os.path.splitext(f.name)[1]
                            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                                tmp.write(f.read())
                                tmp_path = tmp.name
                                new_file_paths.append(tmp_path)
                                st.session_state.processed_files.add(f.name)

                        metadata_list = process_uploaded_files(new_file_paths, status_callback=update_status)
                        embed_documents_from_metadata(metadata_list)

                        if messages:
                            with st.expander("📋 處理紀錄", expanded=True):
                                for m in messages:
                                    st.markdown(f"- {m}")

            elif file_type == "🧪 實驗數據":
                with st.spinner("📥 開始處理實驗資料..."):
                    for f in file_info["files"]:  # 每次都處理，不跳過
                        temp_path = os.path.join("uploaded", f.name)
                        os.makedirs("uploaded", exist_ok=True)
                        with open(temp_path, "wb") as tmp:
                            tmp.write(f.read())

                        result_df, txt_paths = export_new_experiments_to_txt(
                            excel_path=temp_path,
                            output_dir=EXPERIMENT_DIR
                        )

                        

                    embed_experiment_txt_batch(txt_paths)
                    st.success(f"✅ 已處理：{f.name}，共 {len(txt_paths)} 筆")
                    st.dataframe(result_df)

if TAB_FLAGS["tab_4_perplexity_search"]:
    with tab_dict["tab3"]:
        st.subheader("🔍 功能 2：使用關鍵字搜尋外部文獻並下載 PDF")
        q2 = st.text_area("請輸入查詢問題（將自動萃取關鍵字）：", height=100, key="search_pdf")
        if st.button("執行查詢與下載", key="downloadbtn"):
            with st.spinner("🔍 查詢並下載中..."):
                pdfs = search_and_download_only(q2, top_k=5, storage_dir="data/paper")
                if pdfs:
                    st.success(f"✅ 共下載 {len(pdfs)} 篇 PDF")
                    for path in pdfs:
                        st.markdown(f"- `{os.path.basename(path)}`")
                else:
                    st.warning("⚠️ 未成功下載任何 PDF")
if TAB_FLAGS["tab_3_search_pdf"]:
    with tab_dict["tab3"]:
        if "processed_files" not in st.session_state:
            st.session_state.processed_files = set()


        file_info = select_files_paper_mode() 
        print(file_info)
        if file_info["files"]:
            file_type = file_info["type"] 
            
            st.markdown(f"🔎 資料類型：{file_type}")
            for f in file_info["files"]:
                st.write(f"📄 {f.name}")

            new_files = [f for f in file_info["files"] if f.name not in st.session_state.processed_files]

            if not new_files:
                st.info("✅ 所有檔案都已處理過，不重複處理。")
            else:
                with st.spinner("📥 處理論文資料中..."):
                    new_file_paths = []
                    messages = []

                    def update_status(msg):
                        messages.append(msg)

                    for f in new_files:
                        suffix = os.path.splitext(f.name)[1]
                        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                            tmp.write(f.read())
                            tmp_path = tmp.name
                            new_file_paths.append(tmp_path)
                            st.session_state.processed_files.add(f.name)

                    metadata_list = process_uploaded_files(new_file_paths, status_callback=update_status)
                    embed_documents_from_metadata(metadata_list)

                    if messages:
                        with st.expander("📋 處理紀錄", expanded=True):
                            for m in messages:
                                st.markdown(f"- {m}")

if TAB_FLAGS["tab_4_perplexity_search"]:
    with tab_dict["tab4"]:
        st.subheader("🌐 功能 4：使用 Perplexity 搜尋文獻")
        q4 = st.text_area("請輸入要搜尋的文獻問題：", height=100, key="search4")
        if st.button("使用 Perplexity 搜尋", key="perplexitybtn"):
            with st.spinner("使用 Perplexity 搜尋中..."):
                result = ask_perplexity(q4)
                if result["success"]:
                    st.markdown("### 🤖 文獻摘要")
                    st.write(result["answer"])

                    st.markdown("### 🔗 引用來源")
                    links = format_references_block(result["answer"])
                    if links:
                        for link in links:
                            st.markdown(f"- {link}")
                    else:
                        st.info("未偵測到結構化 Reference 區塊。")
                else:
                    st.error("❌ 搜尋失敗：" + result["error"])

if TAB_FLAGS["tab_5_research_assitant"]:
    with tab_dict["tab5"]   :
        st.subheader("📘 功能 1：利用知識庫回答問題")
        st.markdown("### ⚙️ 選擇回答模式")
        answer_mode = st.radio(
            "請選擇系統回答問題的方式：",
            options=[
                "僅嚴謹文獻溯源",
                "允許延伸與推論",
                "納入實驗資料，進行推論與建議"
            ],
            index=0,
            key="mode_selector"
        )

        q1 = st.text_area("請輸入研究問題：", height=100, key="search_2")

        if st.button("由知識庫回答", key="knowledgebtn"):
            with st.spinner("查詢知識庫中..."):
                result = agent_answer(q1, mode = answer_mode)
                st.session_state["proposal_chunks"] = result.get("chunks", [])
                st.session_state["proposal_answer"] = result.get("answer")
                st.session_state["result"] = result
        
            st.markdown("### 🤖 回答")
            st.markdown(st.session_state["proposal_answer"])




            # 📚 引用資料（支援原始或改寫皆可）
            st.markdown("### 📚 引用資料")
            for i, citation in enumerate(st.session_state.get("result", {}).get("citations", []), start=1):
                title = citation.get("title", "未知")
                page = citation.get("page", "?")
                snippet = citation.get("snippet", "...")
                st.markdown(f"**[{i}]** {title} | 頁碼：{page} | 段落開頭：{snippet}")